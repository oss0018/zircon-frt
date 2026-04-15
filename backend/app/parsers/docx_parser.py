from app.parsers.base import BaseParser


class DocxParser(BaseParser):
    """Extract text from DOCX files using python-docx."""

    def parse(self, filepath: str) -> str:
        try:
            import docx
            doc = docx.Document(filepath)
            parts = [para.text for para in doc.paragraphs]
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return "\n".join(p for p in parts if p.strip())
        except Exception:
            return ""
